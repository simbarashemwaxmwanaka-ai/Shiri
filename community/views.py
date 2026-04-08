from django.shortcuts import render, redirect, get_object_or_404
from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
from django.contrib.auth import logout as auth_logout
from django.contrib.auth.views import LoginView
import json
from .forms import MemberForm
from .models import Member

class CustomLoginView(LoginView):
    template_name = 'community/login.html'
    
    def form_valid(self, form):
        # Trigger standard login
        response = super().form_valid(form)
        
        # Modern Session Management: Remember Me functionality
        remember_me = self.request.POST.get('remember_me')
        if remember_me:
            # 2 weeks expiry, survives browser close
            self.request.session.set_expiry(1209600)
        else:
            # Clear passing browser session (or timeout configured in settings)
            self.request.session.set_expiry(0)
            
        return response

def custom_logout(request):
    auth_logout(request)
    return redirect('/')

def home_dashboard(request):
    return render(request, 'community/home.html')

def api_members(request):
    query = request.GET.get('q', '')
    from django.db.models.functions import Length
    if query:
        members = Member.objects.filter(name__icontains=query) | \
                  Member.objects.filter(surname__icontains=query) | \
                  Member.objects.filter(member_number__icontains=query) | \
                  Member.objects.filter(id_number__icontains=query)
    else:
        members = Member.objects.all()
    
    # Sort logically by length then alphabetically (sorts "2" before "10")
    members = members.annotate(num_len=Length('member_number')).order_by('num_len', 'member_number')
    
    data = []
    for m in members:
        data.append({
            'index': m.index,
            'member_number': m.member_number,
            'id_number': m.id_number,
            'name': m.name,
            'surname': m.surname,
            'address': m.address,
            'date_of_birth': str(m.date_of_birth),
            'contact': m.contact,
            'sex': m.sex,
            'educational_qualification': m.educational_qualification,
        })
    return JsonResponse({'members': data})

@login_required
def edit_member(request, index):
    member = get_object_or_404(Member, pk=index)
    if request.method == 'POST':
        form = MemberForm(request.POST, instance=member)
        if form.is_valid():
            form.save()
            return redirect('home_dashboard')
    else:
        form = MemberForm(instance=member)
    
    return render(request, 'community/edit_member.html', {'form': form, 'member': member})

@login_required
def delete_member(request, index):
    if request.method == 'POST':
        member = get_object_or_404(Member, pk=index)
        member.delete()
        return JsonResponse({'status': 'success'})
    return JsonResponse({'status': 'error', 'message': 'Invalid request method'}, status=400)

@login_required
def delete_all_members(request):
    if request.method == 'POST':
        count, _ = Member.objects.all().delete()
        return JsonResponse({'status': 'success', 'deleted': count})
    return JsonResponse({'status': 'error', 'message': 'Invalid request method'}, status=400)

@login_required
def reorder_members(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            order_list = data.get('order', [])
            for new_pos, index_val in enumerate(order_list):
                Member.objects.filter(pk=index_val).update(display_order=new_pos)
            return JsonResponse({'status': 'success'})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=400)
    return JsonResponse({'status': 'error', 'message': 'Invalid request method'}, status=400)

@login_required
def add_member(request):
    if request.method == 'POST':
        form = MemberForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('success')
    else:
        form = MemberForm()
    
    return render(request, 'community/add_member.html', {'form': form})

def success(request):
    return render(request, 'community/success.html')

def get_export_queryset(request):
    query = request.GET.get('q', '')
    from django.db.models.functions import Length
    
    if query:
        qs = Member.objects.filter(name__icontains=query) | \
             Member.objects.filter(surname__icontains=query) | \
             Member.objects.filter(member_number__icontains=query) | \
             Member.objects.filter(id_number__icontains=query)
    else:
        qs = Member.objects.all()
        
    return qs.annotate(num_len=Length('member_number')).order_by('num_len', 'member_number')

@login_required
def export_csv(request):
    import csv
    from django.http import HttpResponse
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="members.csv"'

    writer = csv.writer(response)
    writer.writerow(['Number', 'Full Name', 'Address', 'ID Number', 'Date of Birth', 'Contact', 'Sex', 'Educational Qualification'])

    for m in get_export_queryset(request):
        writer.writerow([m.member_number, f"{m.name} {m.surname}", m.address, m.id_number, str(m.date_of_birth), m.contact, m.sex, m.educational_qualification])

    return response

@login_required
def export_excel(request):
    import openpyxl
    from openpyxl.styles import Font
    from django.http import HttpResponse
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="members.xlsx"'

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Members"
    
    columns = ['Number', 'Full Name', 'Address', 'ID Number', 'Date of Birth', 'Contact', 'Sex', 'Educational Qualification']
    ws.append(columns)
    
    # Bolden the headers
    for cell in ws[1]:
        cell.font = Font(bold=True)
    
    # Widen cells for better readability
    ws.column_dimensions['A'].width = 15 # Number
    ws.column_dimensions['B'].width = 30 # Full Name
    ws.column_dimensions['C'].width = 40 # Address
    ws.column_dimensions['D'].width = 20 # ID Number
    ws.column_dimensions['E'].width = 18 # DOB
    ws.column_dimensions['F'].width = 25 # Contact
    ws.column_dimensions['G'].width = 12 # Sex
    ws.column_dimensions['H'].width = 35 # Educational Qualification

    for m in get_export_queryset(request):
        ws.append([m.member_number, f"{m.name} {m.surname}", m.address, m.id_number, str(m.date_of_birth), m.contact, m.sex, m.educational_qualification])

    wb.save(response)
    return response

@login_required
def export_word(request):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    from django.http import HttpResponse
    
    document = Document()
    document.add_heading('Shiri Water Project Members', 0)
    
    members = get_export_queryset(request)
    table = document.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    columns = ['Number', 'Full Name', 'Address', 'ID Number', 'Date of Birth', 'Contact', 'Sex', 'Educational Qualification']
    for i, col in enumerate(columns):
        p = hdr_cells[i].paragraphs[0]
        p.text = ""
        run = p.add_run(col)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
    for m in members:
        row_cells = table.add_row().cells
        row_data = [
            str(m.member_number), 
            f"{m.name} {m.surname}", 
            str(m.address), 
            str(m.id_number), 
            str(m.date_of_birth), 
            str(m.contact), 
            str(m.sex), 
            str(m.educational_qualification)
        ]
        
        for i, val in enumerate(row_data):
            p = row_cells[i].paragraphs[0]
            p.text = val
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    f = io.BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename="members.docx"'
    response['Content-Length'] = length
    return response

@login_required
def export_pdf(request):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import landscape, letter
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
    import io
    from django.http import HttpResponse

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="members.pdf"'

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter))
    elements = []

    columns = ['No', 'Full Name', 'Address', 'ID Number', 'DOB', 'Contact', 'Sex', 'Education']
    data = [columns]

    for m in get_export_queryset(request):
        data.append([
            str(m.member_number), 
            f"{m.name} {m.surname}", 
            str(m.address), 
            str(m.id_number), 
            str(m.date_of_birth), 
            str(m.contact), 
            str(m.sex), 
            str(m.educational_qualification)
        ])

    table = Table(data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.lightgrey),
        ('GRID', (0,0), (-1,-1), 1, colors.black),
        ('WORDWRAP', (0,0), (-1,-1), True)
    ]))
    
    elements.append(table)
    doc.build(elements)
    
    pdf = buffer.getvalue()
    buffer.close()
    
    response.write(pdf)
    return response

@login_required
def import_excel(request):
    import openpyxl
    from django.http import JsonResponse
    import datetime
    
    if request.method == 'POST' and request.FILES.get('file'):
        excel_file = request.FILES['file']
        
        try:
            wb = openpyxl.load_workbook(excel_file, data_only=True)
            sheet = wb.active
            
            # --- Step 1: Map headers using exact lookup table ---
            # Keys here are the EXACT headings in the Excel file (lowercased + stripped)
            HEADER_MAP = {
                'coo no:':                    'member_number',
                'coo no':                     'member_number',
                'no':                         'member_number',
                'number':                     'member_number',
                'member number':              'member_number',
                'full name':                  'full_name',
                'name':                       'full_name',
                'address':                    'address',
                'id no:':                     'id_number',
                'id no':                      'id_number',
                'id number':                  'id_number',
                'date of birth':              'date_of_birth',
                'dob':                        'date_of_birth',
                'contact details':            'contact',
                'contact':                    'contact',
                'phone':                      'contact',
                'sex':                        'sex',
                'gender':                     'sex',
                'educational qualification':  'educational_qualification',
                'education':                  'educational_qualification',
            }

            headers = {}
            for col_idx, cell in enumerate(sheet[1], 1):
                if cell.value:
                    val = str(cell.value).strip().lower()
                    if val in HEADER_MAP:
                        db_field = HEADER_MAP[val]
                        if db_field not in headers:  # First match wins
                            headers[db_field] = col_idx

            # --- Step 2: Helper defined OUTSIDE the loop to avoid closure bug ---
            def parse_date(val):
                """Convert any common date string format to YYYY-MM-DD."""
                if not val:
                    return '2000-01-01'
                import datetime as dt
                # Already a date/datetime object (native Excel date)
                if isinstance(val, (dt.datetime, dt.date)):
                    return val.strftime('%Y-%m-%d')
                s = str(val).strip()
                # Try all common formats
                for fmt in ('%d/%m/%Y', '%d/%m/%y', '%Y-%m-%d', '%d-%m-%Y',
                            '%d-%m-%y', '%m/%d/%Y', '%m/%d/%y', '%Y/%m/%d'):
                    try:
                        return dt.datetime.strptime(s, fmt).strftime('%Y-%m-%d')
                    except ValueError:
                        pass
                # Last resort: return placeholder so record still saves
                return '2000-01-01'

            def get_val(row, key):
                if key in headers:
                    idx = headers[key] - 1
                    if idx < len(row):
                        v = row[idx]
                        if v is None:
                            return ''
                        # Handle Excel native datetime objects
                        if isinstance(v, (datetime.datetime, datetime.date)):
                            return v.strftime('%Y-%m-%d')
                        return str(v).strip() if str(v).lower() != 'none' else ''
                return ''

            imported_count = 0
            skipped_count = 0

            # --- Step 3: Process each data row ---
            for row in sheet.iter_rows(min_row=2, values_only=True):
                member_num = get_val(row, 'member_number')
                id_num     = get_val(row, 'id_number')
                full_name  = get_val(row, 'full_name')
                address    = get_val(row, 'address')
                dob_raw    = row[headers['date_of_birth'] - 1] if 'date_of_birth' in headers else None
                dob        = parse_date(dob_raw)
                contact    = get_val(row, 'contact')
                sex        = get_val(row, 'sex')
                edu        = get_val(row, 'educational_qualification')

                # Skip completely empty rows
                if not member_num and not full_name:
                    continue

                # Split "Full Name" into name + surname
                parts   = full_name.split(' ', 1)
                name    = parts[0]
                surname = parts[1] if len(parts) > 1 else ''

                # Skip duplicates
                if (member_num and Member.objects.filter(member_number=member_num).exists()) or \
                   (id_num and Member.objects.filter(id_number=id_num).exists()):
                    skipped_count += 1
                    continue

                Member.objects.create(
                    member_number=member_num,
                    id_number=id_num,
                    name=name,
                    surname=surname,
                    address=address,
                    date_of_birth=dob or '2000-01-01',
                    contact=contact,
                    sex=sex,
                    educational_qualification=edu
                )
                imported_count += 1

            return JsonResponse({
                'status': 'success',
                'imported': imported_count,
                'skipped': skipped_count,
                'headers_found': list(headers.keys()),
                'raw_headers': [str(cell.value) for cell in sheet[1] if cell.value]
            })

        except Exception as e:
            import traceback
            return JsonResponse({'status': 'error', 'message': str(e), 'detail': traceback.format_exc()}, status=400)

    return JsonResponse({'status': 'error', 'message': 'Invalid request or missing file payload'}, status=400)

